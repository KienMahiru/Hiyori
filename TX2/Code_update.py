import tkinter as tk
from tkinter import END, Scrollbar, Text, messagebox,simpledialog
import matplotlib.pyplot as plt
import pandas as pd
from numpy import array
import numpy as np
import os
from docx import Document
df = pd.read_csv('diemPython.csv', index_col=0, header=0)
in_data = array(df.iloc[:, :])
global categories1,  svtruot ,tongsv, tongsvtruot
categories1 = ['Lớp 1', 'Lớp 2', 'Lớp 3', 'Lớp 4', 'Lớp 5', 'Lớp 6', 'Lớp 7', 'Lớp 8', 'Lớp 9']
svtruot = in_data[:, 10]
sv = in_data[:, 1]
tongsv = np.sum(sv)
tongsvtruot = np.sum(svtruot)
def xuat_file_word(text_widget, file_path):
    # Lấy nội dung từ ô văn bản
    content = text_widget.get("1.0", "end-1c")
    # Tạo một tài liệu Word mới
    doc = Document()
    # Thêm nội dung vào tài liệu
    paragraph = doc.add_paragraph(content)
    # Thêm các hình ảnh vào tài liệu
    charts = ['chart1.png', 'chart.png', 'chart2.png',"bar_chart.png","pie_chart.png","bar_chart1.png","pie_chart1.png"]
    for chart in charts:
        if os.path.exists(chart):
            doc.add_picture(chart)
    # Lưu tài liệu vào file Word
    doc.save(file_path)
def draw_chart():
    # code vẽ biểu đồ
    plt.savefig('chart.png')
    plt.savefig('chart2.png')
    plt.savefig('chart1.png')
    plt.savefig("bar_chart.png")
    plt.savefig("pie_chart.png")
    plt.savefig("bar_chart1.png")
    plt.savefig("pie_chart1.png")
def xuat_word():
    file_path = "File_sv.docx"
    xuat_file_word(result_text, file_path)
    messagebox.showinfo("Thông báo", "Xuất file thành công")


def hienthids():
    result_text.insert(END, "Danh sách sinh viên:\n")
    result_text.insert(END, df.to_string(index=False) + "\n\n")
def draw_bar_chart(categories, values1, values2):
    total_values = values1 + values2
    fig, ax = plt.subplots(figsize=(12, 6))
    ax.bar(categories, values1, color='blue', label="Sinh viên đạt")
    ax.bar(categories, values2, bottom=values1, color='orange', label="Sinh viên trượt")
    for i, (v1, v2, tv) in enumerate(zip(values1, values2, total_values)):
        ax.text(i, v1 / 2, str(v1), ha='center', va='bottom', color='white', fontweight='bold')
        ax.text(i, v1 + v2 / 2, str(v2), ha='center', va='bottom', color='white', fontweight='bold')
        ax.text(i, v1 + v2, str(tv), ha='center', va='bottom', color='black', fontweight='bold')
    ax.set_title('Biểu đồ số sinh viên trong các lớp')
    ax.set_ylabel('Số sinh viên')
    ax.set_xlabel('Lớp')
    ax.legend(loc='upper right')
    plt.savefig('bar_chart.png')
    plt.show()

def draw_pie_chart(percent_dat, percent_truot):
    labels = ["Sinh viên đạt", "Sinh viên trượt"]
    fig, ax = plt.subplots(figsize=(6, 6))
    ax.pie([percent_dat, percent_truot], labels=labels, autopct='%1.1f%%')
    ax.set_title('Biểu đồ tỷ lệ sinh viên đạt và trượt')
    plt.savefig('pie_chart.png')
    plt.show()

def tongsinhvien():
    result_text.insert(END, f"Tổng số sinh viên đi thi: {tongsv} sinh viên\n")
    tongsvtruot = np.sum(svtruot)
    svdat = tongsv - tongsvtruot
    percent_dat = (svdat / tongsv) * 100
    percent_truot = (tongsvtruot / tongsv) * 100
    result_text.insert(END, f"Tổng số sinh viên qua môn: {svdat} sinh viên chiếm {percent_dat}%\n")
    result_text.insert(END, f"Tổng số sinh viên trượt môn: {tongsvtruot} sinh viên chiếm {percent_truot}%\n")
    values1 = np.sum(in_data[0:9, 2:10], axis=1).flatten()
    values2 = in_data[:, 10]
    draw_bar_chart(categories1, values1, values2)
    draw_pie_chart(percent_dat, percent_truot)
def sosinhvientruot():

    result_text.insert(END, "Tổng số sinh viên trượt môn:")
    result_text.insert(END, str(tongsvtruot) + "\n\n")
    values1 = in_data[:, 10]
    plt.figure(2)
    plt.bar(categories1, values1, color='red', label="Sinh viên trượt")
    plt.title('Biểu đồ số sinh viên trượt của các lớp')
    plt.ylabel('Số sinh viên')
    plt.legend(loc='upper right')
    plt.show()
def sosinhviendat():
    tongsvdat = tongsv - tongsvtruot
    result_text.insert(END, f"Tổng số sinh viên qua môn: {tongsvdat}sinh viên \n")
    values2 = np.sum(in_data[0:9, 2:10], axis=1).flatten()
    plt.figure(3)
    plt.bar(categories1, values2, color='green', label="Sinh viên đạt")
    plt.title('Biểu đồ số sinh viên đạt của các lớp')
    plt.ylabel('Số sinh viên')
    plt.legend(loc='upper right')
    plt.show()
def sinhviendatdiem():
    root = tk.Tk()
    root.geometry("500x200")
    root.withdraw()
    while True:
        # Mở hộp thoại nhập điểm
        user_input = simpledialog.askstring("Nhập điểm:", "Nhập điểm của sinh viên(ví dụ: A):")
        if user_input is None:
            result_text.insert(tk.END, "Bạn không nhập điểm! Vui lòng nhập lại!\n")
            break
        elif user_input.upper() != "A" or user_input.upper() != "B" or user_input.upper() != "C+" or user_input.upper() != "B+" or user_input.upper() != "C" or user_input.upper() != "D+" or user_input.upper() != "D" or user_input.upper() != "F":
            result_text.insert(tk.END, "Điểm không hợp lệ! Vui lòng nhập lại!\n")
            break
        else:
            result_text.insert(tk.END, "Bạn nhập điểm không hợp lệ! Vui lòng nhập lại!\n")
            break
    result_text.insert(tk.END, "Bạn nhập điểm là: " + str(user_input) + "\n")
    global categories1
    categories1 = ['Lớp 1', 'Lớp 2', 'Lớp 3', 'Lớp 4', 'Lớp 5', 'Lớp 6', 'Lớp 7', 'Lớp 8', 'Lớp 9']
    if user_input == "A":
    # xử lý điểm A
        diemA = in_data[:, 3]
        maxa = diemA.max()
        i = np.argmax(diemA)
        result_text.insert(END, '\nLớp có nhiều sinh viên điểm A là lớp {0} có {1} sinh viên\n'.format(in_data[i, 0], maxa))
        values1 = in_data[:, 3]
        plt.figure(4)
        bars = plt.bar(categories1, values1, label="Sinh viên đạt điểm A")
        plt.title('Biểu đồ số sinh viên đạt điểm A của các lớp')
        plt.ylabel('Số sinh viên')
        plt.legend(loc='upper right')
        for bar, value in zip(bars, values1):
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width() / 2, height, value, ha='center', va='bottom')
        plt.show()
    if user_input == "B+":
    # xử lý điểm B+
        diemA = in_data[:, 4]
        maxa = diemA.max()
        i = np.argmax(diemA)
        result_text.insert(END, '\nLớp có nhiều sinh viên điểm B+ là lớp {0} có {1} sinh viên\n'.format(in_data[i, 0], maxa))
        values1 = in_data[:, 4]
        plt.figure(5)
        bars = plt.bar(categories1, values1, label="Sinh viên đạt điểm B")
        plt.title('Biểu đồ số sinh viên đạt điểm B+ của các lớp')
        plt.ylabel('Số sinh viên')
        plt.legend(loc='upper right')
        for bar, value in zip(bars, values1):
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width() / 2, height, value, ha='center', va='bottom')
        plt.show()
    if user_input == "B":
        # xử lý điểm B
        diemA = in_data[:, 5]
        maxa = diemA.max()
        i = np.argmax(diemA)
        result_text.insert(END,
                           '\nLớp có nhiều sinh viên điểm B là lớp {0} có {1} sinh viên\n'.format(in_data[i, 0], maxa))
        values1 = in_data[:, 5]
        plt.figure(6)
        bars = plt.bar(categories1, values1, label="Sinh viên đạt điểm B")
        plt.title('Biểu đồ số sinh viên đạt điểm B của các lớp')
        plt.ylabel('Số sinh viên')
        plt.legend(loc='upper right')
        for bar, value in zip(bars, values1):
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width() / 2, height, value, ha='center', va='bottom')
        plt.show()
    if user_input == "C+":
        # xử lý điểm C+
        diemA = in_data[:, 6]
        maxa = diemA.max()
        i = np.argmax(diemA)
        result_text.insert(END, '\nLớp có nhiều sinh viên điểm C+ là lớp {0} có {1} sinh viên\n'.format(in_data[i, 0], maxa))

        values1 = in_data[:, 6]
        plt.figure(7)
        bars = plt.bar(categories1, values1, label="Sinh viên đạt điểm C+")
        plt.title('Biểu đồ số sinh viên đạt điểm C+ của các lớp')
        plt.ylabel('Số sinh viên')
        plt.legend(loc='upper right')
        for bar, value in zip(bars, values1):
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width() / 2, height, value, ha='center', va='bottom')
        plt.show()
    if user_input == "C":
        # xử lý điểm C
        diemA = in_data[:, 7]
        maxa = diemA.max()
        i = np.argmax(diemA)
        result_text.insert(END, '\nLớp có nhiều sinh viên điểm C là lớp {0} có {1} sinh viên\n'.format(in_data[i, 0], maxa))
        values1 = in_data[:, 7]
        plt.figure(8)
        bars = plt.bar(categories1, values1, label="Sinh viên đạt điểm C")
        plt.title('Biểu đồ số sinh viên đạt điểm C của các lớp')
        plt.ylabel('Số sinh viên')
        plt.legend(loc='upper right')
        for bar, value in zip(bars, values1):
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width() / 2, height, value, ha='center', va='bottom')
        plt.show()
    if user_input == "C":
        # xử lý điểm C
        diemA = in_data[:, 7]
        maxa = diemA.max()
        i = np.argmax(diemA)
        result_text.insert(END, '\nLớp có nhiều sinh viên điểm C là lớp {0} có {1} sinh viên\n'.format(in_data[i, 0], maxa))
        values1 = in_data[:, 7]
        plt.figure(8)
        bars = plt.bar(categories1, values1, label="Sinh viên đạt điểm C")
        plt.title('Biểu đồ số sinh viên đạt điểm C của các lớp')
        plt.ylabel('Số sinh viên')
        plt.legend(loc='upper right')
        for bar, value in zip(bars, values1):
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width() / 2, height, value, ha='center', va='bottom')
        plt.show()
    if user_input == "D+":
        # xử lý điểm D+
        diemA = in_data[:, 8]
        maxa = diemA.max()
        i = np.argmax(diemA)
        result_text.insert(END,
                           '\nLớp có nhiều sinh viên điểm D+ là lớp {0} có {1} sinh viên\n'.format(in_data[i, 0], maxa))
        categories1 = ['Lớp 1', 'Lớp 2', 'Lớp 3', 'Lớp 4', 'Lớp 5', 'Lớp 6', 'Lớp 7', 'Lớp 8', 'Lớp 9']
        values1 = in_data[:, 8]
        plt.figure(9)
        bars = plt.bar(categories1, values1, label="Sinh viên đạt điểm D+")
        plt.title('Biểu đồ số sinh viên đạt điểm D+ của các lớp')
        plt.ylabel('Số sinh viên')
        plt.legend(loc='upper right')
        for bar, value in zip(bars, values1):
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width() / 2, height, value, ha='center', va='bottom')
        plt.show()
    if user_input == "D":
        # xử lý điểm D
        diemA = in_data[:, 9]
        maxa = diemA.max()
        i = np.argmax(diemA)
        result_text.insert(END, '\nLớp có nhiều sinh viên điểm D là lớp {0} có {1} sinh viên\n'.format(in_data[i, 0], maxa))
        values1 = in_data[:, 9]
        plt.figure(10)
        bars = plt.bar(categories1, values1, label="Sinh viên đạt điểm D")
        plt.title('Biểu đồ số sinh viên đạt điểm D của các lớp')
        plt.ylabel('Số sinh viên')
        plt.legend(loc='upper right')
        for bar, value in zip(bars, values1):
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width() / 2, height, value, ha='center', va='bottom')
        plt.show()
    if user_input == "F":
        # xử lý điểm F
        diemA = in_data[:, 9]
        maxa = diemA.max()
        i = np.argmax(diemA)
        result_text.insert(END,'\nLớp có nhiều sinh viên điểm D là lớp {0} có {1} sinh viên\n'.format(in_data[i, 0], maxa))
        values1 = in_data[:, 9]
        plt.figure(11)
        bars = plt.bar(categories1, values1, label="Sinh viên đạt điểm D")
        plt.title('Biểu đồ số sinh viên đạt điểm D của các lớp')
        plt.ylabel('Số sinh viên')
        plt.legend(loc='upper right')
        for bar, value in zip(bars, values1):
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width() / 2, height, value, ha='center', va='bottom')
        plt.show()
import matplotlib.pyplot as plt

def draw_bar_chart_pho(labels, grade_counts):
    colors = ['red', 'green', 'yellow', 'blue', 'orange', 'purple', 'pink', 'brown']
    fig, ax = plt.subplots(figsize=(8, 6))
    ax.bar(labels, grade_counts, color=colors)
    ax.set_xlabel('Điểm')
    ax.set_ylabel('Số lượng sinh viên')
    ax.set_title('Phổ điểm')
    for i, v in enumerate(grade_counts):
        ax.annotate(str(v), (i, v), ha='center', va='bottom')
    plt.savefig('bar_chart1.png')
    plt.show()

def draw_pie_chart_pho(labels, grade_counts):
    colors = ['red', 'green', 'yellow', 'blue', 'orange', 'purple', 'pink', 'brown']
    fig, ax = plt.subplots(figsize=(6, 6))
    ax.pie(grade_counts, labels=labels, colors=colors, autopct='%1.1f%%')
    ax.axis('equal')
    ax.set_title('Phổ điểm')
    plt.savefig('pie_chart1.png')
    plt.show()

def phodiem():
    labels = ["A", "B+", "B", "C+", "C", "D+", "D", "F"]
    diems = in_data[:, 3:11]
    grade_counts = diems.sum(axis=0)
    draw_bar_chart_pho(labels, grade_counts)
    draw_pie_chart_pho(labels, grade_counts)

# Rest of your code...
def Xoa():
    result_text.delete(1.0, END)
    plt.close('all')

# Tạo cửa sổ giao diện
window = tk.Tk()
window.title("Phân tích điểm sinh viên ")
# Tạo các nút chức năng
htdanhsach_button = tk.Button(window, text="Hiển thị danh sách", command=hienthids)
tongsinhvien_button = tk.Button(window, text="Tổng sinh viên", command=tongsinhvien)
sinhvientruot_button = tk.Button(window, text="Sinh viên trượt", command=sosinhvientruot)
sinhviendat_button = tk.Button(window, text="Sinh viên đạt", command=sosinhviendat)

sinhvienA_button = tk.Button(window, text="Sinh viên đạt điểm", command=sinhviendatdiem)
phodiem_button = tk.Button(window, text="Phổ điểm", command=phodiem)
reset_button = tk.Button(window, text="Xóa", command=Xoa)
xuat_word_button = tk.Button(window, text="Xuất file word", command=xuat_word)
# Tạo ô văn bản để hiển thị kết quả
result_text = Text(window, wrap="word", height=20, width=200)
result_text_scrollbar = Scrollbar(window, command=result_text.yview)
result_text.config(yscrollcommand=result_text_scrollbar.set)
# Đặt vị trí các phần tử trên cửa sổ
htdanhsach_button.grid(row=0, column=0, padx=10, pady=10, sticky="ew")
tongsinhvien_button.grid(row=1, column=0, padx=10, pady=10, sticky="ew")
sinhvientruot_button.grid(row=2, column=0, padx=10, pady=10, sticky="ew")
sinhviendat_button.grid(row=3, column=0, padx=10, pady=10, sticky="ew")
sinhvienA_button.grid(row=4, column=0, padx=10, pady=10, sticky="ew")
phodiem_button.grid(row=5, column=0, padx=10, pady=10, sticky="ew")
reset_button.grid(row=6, column=0, padx=10, pady=10, sticky="ew")
result_text.grid(row=0, column=1, rowspan=7, padx=10, pady=10, sticky="nsew")
result_text_scrollbar.grid(row=0, column=2, rowspan=7, sticky="ns")
xuat_word_button.grid(row=7, column=0, padx=10, pady=10, sticky="ew")
# Chạy ứng dụng
window.mainloop()